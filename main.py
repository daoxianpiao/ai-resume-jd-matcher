from __future__ import annotations

import json
import os
import sqlite3
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Literal, Optional

from dotenv import load_dotenv
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi import Response
from fastapi.responses import HTMLResponse
from fastapi.templating import Jinja2Templates
from openai import OpenAI, OpenAIError
from pydantic import BaseModel, Field
from pypdf import PdfReader


load_dotenv()

app = FastAPI(title="AI Resume JD Matcher")
templates = Jinja2Templates(directory="templates")

DATABASE_PATH = Path(__file__).resolve().parent / "data" / "app.db"
MAX_UPLOAD_BYTES = 5 * 1024 * 1024
SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
PLACEHOLDER_API_KEYS = {"", "your_openai_api_key_here"}


class ScoreItem(BaseModel):
    name: str
    score: int = Field(ge=0, le=100)
    reason: str


class GapItem(BaseModel):
    skill: str
    severity: Literal["high", "medium", "low"]
    evidence: str
    action: str


class ProjectSuggestion(BaseModel):
    title: str
    why: str
    features: list[str]


class ResumeRewriteSuggestion(BaseModel):
    before_issue: str
    rewrite_tip: str


class LearningTaskUpdate(BaseModel):
    is_completed: bool


class MatchAnalysis(BaseModel):
    target_role: str
    overall_score: int = Field(ge=0, le=100)
    summary: str
    strengths: list[str]
    risks: list[str]
    score_breakdown: list[ScoreItem]
    missing_keywords: list[str]
    gaps: list[GapItem]
    learning_plan_30_days: list[str]
    project_suggestions: list[ProjectSuggestion]
    resume_rewrite_suggestions: list[ResumeRewriteSuggestion]
    interview_talking_points: list[str]


def get_db_connection() -> sqlite3.Connection:
    DATABASE_PATH.parent.mkdir(parents=True, exist_ok=True)
    connection = sqlite3.connect(DATABASE_PATH)
    connection.row_factory = sqlite3.Row
    return connection


def init_db() -> None:
    with get_db_connection() as connection:
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS analyses (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                created_at TEXT NOT NULL,
                analysis_mode TEXT NOT NULL,
                target_role TEXT NOT NULL,
                overall_score INTEGER NOT NULL,
                resume_preview TEXT NOT NULL,
                job_preview TEXT NOT NULL,
                resume TEXT NOT NULL,
                job_description TEXT NOT NULL,
                result_json TEXT NOT NULL
            )
            """
        )
        connection.execute(
            """
            CREATE TABLE IF NOT EXISTS learning_tasks (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                analysis_id INTEGER NOT NULL,
                title TEXT NOT NULL,
                position INTEGER NOT NULL,
                is_completed INTEGER NOT NULL DEFAULT 0,
                created_at TEXT NOT NULL,
                completed_at TEXT,
                FOREIGN KEY (analysis_id) REFERENCES analyses(id) ON DELETE CASCADE
            )
            """
        )
        connection.execute(
            """
            CREATE INDEX IF NOT EXISTS idx_learning_tasks_analysis_id
            ON learning_tasks (analysis_id, position)
            """
        )


def preview_text(text: str, limit: int = 80) -> str:
    compact = " ".join(text.split())
    if len(compact) <= limit:
        return compact
    return compact[:limit].rstrip() + "..."


def save_analysis_history(
    resume: str,
    job_description: str,
    analysis_mode: str,
    result_data: dict,
) -> int:
    init_db()
    created_at = datetime.now().astimezone().isoformat(timespec="seconds")
    with get_db_connection() as connection:
        cursor = connection.execute(
            """
            INSERT INTO analyses (
                created_at,
                analysis_mode,
                target_role,
                overall_score,
                resume_preview,
                job_preview,
                resume,
                job_description,
                result_json
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                created_at,
                analysis_mode,
                result_data["target_role"],
                result_data["overall_score"],
                preview_text(resume),
                preview_text(job_description),
                resume,
                job_description,
                json.dumps(result_data, ensure_ascii=False),
            ),
        )
        return int(cursor.lastrowid)


def learning_task_from_row(row: sqlite3.Row) -> dict:
    return {
        "id": row["id"],
        "analysis_id": row["analysis_id"],
        "title": row["title"],
        "position": row["position"],
        "is_completed": bool(row["is_completed"]),
        "created_at": row["created_at"],
        "completed_at": row["completed_at"],
    }


def create_learning_tasks(analysis_id: int, items: list[str]) -> list[dict]:
    init_db()
    now = datetime.now().astimezone().isoformat(timespec="seconds")
    clean_items = [item.strip() for item in items if item.strip()]
    if not clean_items:
        return []

    with get_db_connection() as connection:
        connection.executemany(
            """
            INSERT INTO learning_tasks (
                analysis_id,
                title,
                position,
                created_at
            )
            VALUES (?, ?, ?, ?)
            """,
            [
                (analysis_id, title, index + 1, now)
                for index, title in enumerate(clean_items)
            ],
        )

    return list_learning_tasks(analysis_id=analysis_id)


def ensure_learning_tasks_for_analysis(analysis_id: int, result: dict) -> list[dict]:
    tasks = list_learning_tasks(analysis_id=analysis_id)
    if tasks:
        return tasks
    return create_learning_tasks(
        analysis_id=analysis_id,
        items=result.get("learning_plan_30_days", []),
    )


def list_learning_tasks(analysis_id: Optional[int] = None) -> list[dict]:
    init_db()
    query = """
        SELECT
            id,
            analysis_id,
            title,
            position,
            is_completed,
            created_at,
            completed_at
        FROM learning_tasks
    """
    params: tuple[int, ...] = ()
    if analysis_id is not None:
        query += " WHERE analysis_id = ?"
        params = (analysis_id,)
    query += " ORDER BY analysis_id DESC, position ASC, id ASC"

    with get_db_connection() as connection:
        rows = connection.execute(query, params).fetchall()

    return [learning_task_from_row(row) for row in rows]


def get_learning_task(task_id: int) -> dict:
    init_db()
    with get_db_connection() as connection:
        row = connection.execute(
            """
            SELECT
                id,
                analysis_id,
                title,
                position,
                is_completed,
                created_at,
                completed_at
            FROM learning_tasks
            WHERE id = ?
            """,
            (task_id,),
        ).fetchone()

    if row is None:
        raise HTTPException(status_code=404, detail="Learning task not found.")

    return learning_task_from_row(row)


def update_learning_task(task_id: int, is_completed: bool) -> dict:
    init_db()
    completed_at = (
        datetime.now().astimezone().isoformat(timespec="seconds")
        if is_completed
        else None
    )
    with get_db_connection() as connection:
        cursor = connection.execute(
            """
            UPDATE learning_tasks
            SET is_completed = ?, completed_at = ?
            WHERE id = ?
            """,
            (1 if is_completed else 0, completed_at, task_id),
        )

    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Learning task not found.")

    return get_learning_task(task_id)


def delete_learning_task(task_id: int) -> dict:
    init_db()
    with get_db_connection() as connection:
        cursor = connection.execute("DELETE FROM learning_tasks WHERE id = ?", (task_id,))

    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Learning task not found.")

    return {"deleted": True, "id": task_id}


def list_analysis_history(limit: int = 20) -> list[dict]:
    init_db()
    safe_limit = max(1, min(limit, 50))
    with get_db_connection() as connection:
        rows = connection.execute(
            """
            SELECT
                id,
                created_at,
                analysis_mode,
                target_role,
                overall_score,
                resume_preview,
                job_preview,
                result_json,
                (
                    SELECT COUNT(*)
                    FROM learning_tasks
                    WHERE learning_tasks.analysis_id = analyses.id
                ) AS task_total,
                (
                    SELECT COUNT(*)
                    FROM learning_tasks
                    WHERE learning_tasks.analysis_id = analyses.id
                        AND learning_tasks.is_completed = 1
                ) AS task_completed
            FROM analyses
            ORDER BY id DESC
            LIMIT ?
            """,
            (safe_limit,),
        ).fetchall()

    items = []
    for row in rows:
        item = dict(row)
        result = json.loads(item.pop("result_json"))
        item["missing_keywords"] = result.get("missing_keywords", [])[:5]
        items.append(item)
    return items


def get_analysis_history_item(analysis_id: int) -> dict:
    init_db()
    with get_db_connection() as connection:
        row = connection.execute(
            """
            SELECT
                id,
                created_at,
                analysis_mode,
                resume,
                job_description,
                result_json
            FROM analyses
            WHERE id = ?
            """,
            (analysis_id,),
        ).fetchone()

    if row is None:
        raise HTTPException(status_code=404, detail="Analysis history item not found.")

    result = json.loads(row["result_json"])
    result["analysis_id"] = row["id"]
    result["analysis_mode"] = row["analysis_mode"]
    result["learning_tasks"] = ensure_learning_tasks_for_analysis(row["id"], result)
    return {
        "id": row["id"],
        "created_at": row["created_at"],
        "analysis_mode": row["analysis_mode"],
        "resume": row["resume"],
        "job_description": row["job_description"],
        "result": result,
    }


def delete_analysis_history_item(analysis_id: int) -> dict:
    init_db()
    with get_db_connection() as connection:
        connection.execute("DELETE FROM learning_tasks WHERE analysis_id = ?", (analysis_id,))
        cursor = connection.execute("DELETE FROM analyses WHERE id = ?", (analysis_id,))

    if cursor.rowcount == 0:
        raise HTTPException(status_code=404, detail="Analysis history item not found.")

    return {"deleted": True, "id": analysis_id}


def clear_analysis_history() -> dict:
    init_db()
    with get_db_connection() as connection:
        connection.execute("DELETE FROM learning_tasks")
        cursor = connection.execute("DELETE FROM analyses")
    return {"deleted": cursor.rowcount}


def has_openai_api_key() -> bool:
    api_key = (os.getenv("OPENAI_API_KEY") or "").strip()
    return api_key not in PLACEHOLDER_API_KEYS


def get_client() -> OpenAI:
    if not has_openai_api_key():
        raise HTTPException(
            status_code=400,
            detail="OPENAI_API_KEY is missing. Create a .env file from .env.example first.",
        )
    return OpenAI()


def build_prompt(resume: str, job_description: str) -> str:
    return f"""
请分析下面的候选人简历和岗位 JD，判断候选人对岗位的匹配程度。

要求：
1. 所有结论必须基于简历和 JD 文本，不要编造经历。
2. 如果简历没有证据支持，要明确指出缺口。
3. 建议要偏向 AI 应用开发、AI Agent 应用开发、RAG、LLM API、FastAPI、Python、基础前端这些求职方向。
4. 输出中文，语气直接、具体、适合本科生执行。
5. 学习计划要能在 30 天内开始执行，不要空泛。

候选人简历：
{resume}

岗位 JD：
{job_description}
""".strip()


def decode_text_file(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="Text file encoding is not supported.")


def extract_pdf_text(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())
    return "\n\n".join(pages)


def extract_docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def extract_upload_text(upload: UploadFile) -> str:
    filename = upload.filename or ""
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Use: {supported}")

    content = upload.file.read()
    if not content:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail="File is too large. Max size is 5 MB.")

    try:
        if extension in {".txt", ".md"}:
            text = decode_text_file(content)
        elif extension == ".pdf":
            text = extract_pdf_text(content)
        else:
            text = extract_docx_text(content)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to extract text: {exc}") from exc

    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=400,
            detail="No text could be extracted. Scanned PDFs need OCR, which is not enabled yet.",
        )

    return text


def keyword_present(text: str, keyword: str) -> bool:
    return keyword.lower() in text.lower()


def generate_demo_analysis(resume: str, job_description: str) -> MatchAnalysis:
    resume_text = resume.lower()
    jd_text = job_description.lower()
    keywords = [
        "Python",
        "FastAPI",
        "OpenAI API",
        "RAG",
        "Embedding",
        "Agent",
        "Docker",
        "MySQL",
        "Redis",
        "React",
        "Next.js",
        "Linux",
    ]
    missing_keywords = [
        keyword
        for keyword in keywords
        if keyword.lower() in jd_text and keyword.lower() not in resume_text
    ][:8]

    python_score = 70 if keyword_present(resume, "Python") else 45
    ai_score = 62 if any(keyword_present(resume, item) for item in ["RAG", "OpenAI", "LLM"]) else 42
    backend_score = 58 if any(keyword_present(resume, item) for item in ["FastAPI", "Flask", "Django"]) else 35
    frontend_score = 40 if any(keyword_present(resume, item) for item in ["React", "Vue", "HTML", "CSS"]) else 22
    project_score = 52 if any(keyword_present(resume, item) for item in ["项目", "project", "GitHub"]) else 30
    overall_score = round((python_score + ai_score + backend_score + frontend_score + project_score) / 5)

    return MatchAnalysis(
        target_role="AI 应用开发 / Vibe Coding 研发工程师",
        overall_score=overall_score,
        summary=(
            "当前是演示模式：结果由本地规则生成，不代表真实模型判断。"
            "从文本看，你适合先走 Python + FastAPI + OpenAI API + RAG 的项目路线，"
            "用作品补足前端和工程化短板。"
        ),
        strengths=[
            "人工智能专业背景有利于理解 LLM、RAG、Prompt、模型评估等概念。",
            "已经具备 Python 基础，适合从 FastAPI 后端切入 AI 应用开发。",
            "对 AI Native / Vibe Coding 方向兴趣明确，项目选题可以直接服务求职。",
        ],
        risks=[
            "前端基础偏弱，短期内需要补 HTML、CSS、JavaScript 的最小闭环能力。",
            "如果缺少可运行项目，简历很难证明你能独立完成业务落地。",
            "数据库、Docker、部署、日志这些工程化能力需要通过项目逐步补齐。",
        ],
        score_breakdown=[
            ScoreItem(name="Python 基础", score=python_score, reason="Python 是你的主语言，但还需要通过接口、文件处理、数据库项目加强熟练度。"),
            ScoreItem(name="AI 应用理解", score=ai_score, reason="适合继续学习 LLM API、结构化输出、RAG 和工具调用。"),
            ScoreItem(name="后端工程", score=backend_score, reason="需要把 FastAPI、REST API、错误处理、数据库保存做成完整项目。"),
            ScoreItem(name="前端交互", score=frontend_score, reason="先掌握表单、上传、列表、异步请求和结果渲染，不急着上复杂框架。"),
            ScoreItem(name="项目闭环", score=project_score, reason="求职最需要可展示项目、README、截图、部署地址和面试讲解。"),
        ],
        missing_keywords=missing_keywords or ["FastAPI", "RAG", "Docker", "数据库", "部署", "结构化输出"],
        gaps=[
            GapItem(
                skill="前端基础",
                severity="high",
                evidence="你之前说明自己不会前端，而岗位通常要求能完成基础页面和联调。",
                action="先学 HTML/CSS/JS 表单、fetch、DOM 渲染；本项目页面就是第一份练习材料。",
            ),
            GapItem(
                skill="后端与数据库",
                severity="medium",
                evidence="岗位常要求 MySQL、Redis、接口开发和部署经验。",
                action="下一步给本项目加入 SQLite 历史记录，再升级 MySQL/Docker。",
            ),
            GapItem(
                skill="AI 应用作品",
                severity="medium",
                evidence="AI 应用岗位更看重能运行的 RAG、Agent、LLM API 项目。",
                action="完成简历匹配助手、知识库 RAG、轻量 Agent 三个项目，形成作品集。",
            ),
        ],
        learning_plan_30_days=[
            "第 1 周：读懂本项目 main.py，掌握 FastAPI 路由、Form、UploadFile、Pydantic。",
            "第 2 周：补 HTML/CSS/JS，能独立修改页面布局、按钮事件和 fetch 请求。",
            "第 3 周：加入 SQLite 保存分析历史，理解表结构、增删查改和接口返回。",
            "第 4 周：接入真实 OpenAI API 或兼容模型，补 README、截图和部署说明。",
        ],
        project_suggestions=[
            ProjectSuggestion(
                title="AI 简历 JD 匹配助手",
                why="直接服务你的求职，也覆盖文件解析、结构化输出和前后端联调。",
                features=["上传简历/JD", "匹配度评分", "能力缺口分析", "历史记录保存"],
            ),
            ProjectSuggestion(
                title="个人知识库 RAG 问答系统",
                why="RAG 是 AI 应用开发岗位最常见关键词之一。",
                features=["文档切分", "Embedding", "向量检索", "回答引用来源"],
            ),
        ],
        resume_rewrite_suggestions=[
            ResumeRewriteSuggestion(
                before_issue="只写“会 Python / 学过 AI 课程”说服力不足。",
                rewrite_tip="改成“基于 FastAPI + OpenAI API 开发简历 JD 匹配助手，支持文件解析、结构化输出和结果可视化”。",
            ),
            ResumeRewriteSuggestion(
                before_issue="没有体现 AI Coding / Vibe Coding 方法。",
                rewrite_tip="补充“使用 Codex 辅助需求拆解、接口实现、错误定位和 README 编写”。",
            ),
        ],
        interview_talking_points=[
            "我选择先用 FastAPI 做后端，因为 AI 应用岗位更需要快速把模型能力封装成业务接口。",
            "这个项目把简历/JD 文件解析成文本，再调用模型输出结构化结果，前端按字段渲染。",
            "我目前的短板是前端和部署，所以项目规划里会继续加入数据库、Docker 和部署能力。",
        ],
    )


def analyze_resume(resume: str, job_description: str) -> MatchAnalysis:
    if not has_openai_api_key():
        return generate_demo_analysis(resume=resume, job_description=job_description)

    client = get_client()
    model = os.getenv("OPENAI_MODEL", "gpt-5.4-mini")

    try:
        response = client.responses.parse(
            model=model,
            input=[
                {
                    "role": "system",
                    "content": (
                        "你是 AI 应用开发方向的技术招聘分析师，擅长把岗位要求拆成"
                        "工程能力、AI 能力、项目能力和简历表达建议。"
                    ),
                },
                {"role": "user", "content": build_prompt(resume, job_description)},
            ],
            text_format=MatchAnalysis,
            reasoning={"effort": "low"},
        )
    except OpenAIError as exc:
        raise HTTPException(status_code=502, detail=f"OpenAI API error: {exc}") from exc

    if response.output_parsed is None:
        raise HTTPException(status_code=502, detail="OpenAI did not return a parsed result.")

    return response.output_parsed


@app.get("/", response_class=HTMLResponse)
def index(request: Request):
    return templates.TemplateResponse(request, "index.html")


@app.get("/favicon.ico", include_in_schema=False)
def favicon():
    return Response(status_code=204)


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/api/analyze")
def analyze(
    resume: str = Form(..., min_length=30),
    job_description: str = Form(..., min_length=30),
):
    result = analyze_resume(resume=resume, job_description=job_description)
    analysis_mode = "openai" if has_openai_api_key() else "demo"
    data = result.model_dump()
    data["analysis_mode"] = analysis_mode
    data["analysis_id"] = save_analysis_history(
        resume=resume,
        job_description=job_description,
        analysis_mode=analysis_mode,
        result_data=data,
    )
    data["learning_tasks"] = create_learning_tasks(
        analysis_id=data["analysis_id"],
        items=data["learning_plan_30_days"],
    )
    return data


@app.post("/api/extract-file")
def extract_file(file: UploadFile = File(...)):
    text = extract_upload_text(file)
    return {"filename": file.filename, "text": text, "characters": len(text)}


@app.get("/api/history")
def history(limit: int = 20):
    return {"items": list_analysis_history(limit=limit)}


@app.get("/api/history/{analysis_id}")
def history_item(analysis_id: int):
    return get_analysis_history_item(analysis_id)


@app.delete("/api/history/{analysis_id}")
def delete_history_item(analysis_id: int):
    return delete_analysis_history_item(analysis_id)


@app.delete("/api/history")
def clear_history():
    return clear_analysis_history()


@app.get("/api/learning-tasks")
def learning_tasks(analysis_id: Optional[int] = None):
    return {"items": list_learning_tasks(analysis_id=analysis_id)}


@app.patch("/api/learning-tasks/{task_id}")
def patch_learning_task(task_id: int, payload: LearningTaskUpdate):
    return update_learning_task(task_id=task_id, is_completed=payload.is_completed)


@app.delete("/api/learning-tasks/{task_id}")
def remove_learning_task(task_id: int):
    return delete_learning_task(task_id)
